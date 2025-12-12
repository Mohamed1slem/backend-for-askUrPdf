import os
import io
import fitz  # PyMuPDF
import docx
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
import re
from config import RAW_FOLDER, CLEAN_TEXT_FOLDER, CHUNKS_OUTPUT

# ---------- Tesseract Setup ----------
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
os.environ["TESSDATA_PREFIX"] = r"C:\Program Files\Tesseract-OCR\tessdata"

# ---------- Text Cleaning ----------
def clean_text(text):
    """Clean OCR/text extraction output for better chunking."""
    # Remove spaces inside words (common with OCR)
    text = re.sub(r'(?<=\w) (?=\w)', '', text)
    # Normalize multiple newlines
    text = re.sub(r'\n+', '\n', text)
    # Remove unwanted characters (keep letters, numbers, punctuation, Arabic, French accents)
    text = re.sub(r'[^\w\s.,:;!?()\[\]{}\-\'\"ء-يÀ-ÿ]', '', text)
    return text.strip()

# ---------- Text Extraction ----------
def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

# ---------- Image Extraction ----------
def extract_images_from_docx(file_path):
    doc = docx.Document(file_path)
    images = []
    for rel in doc.part._rels:
        rel_obj = doc.part._rels[rel]
        if "image" in rel_obj.target_ref:
            try:
                images.append(rel_obj.target_part.blob)
            except:
                continue
    return images

def extract_images_from_pdf(file_path):
    doc = fitz.open(file_path)
    images = []
    for page in doc:
        for img in page.get_images(full=True):
            xref = img[0]
            base_image = doc.extract_image(xref)
            images.append(base_image["image"])
    return images

# ---------- OCR ----------
def ocr_images(image_bytes_list):
    """Perform OCR on images with French + Arabic languages and preprocessing."""
    texts = []
    for i, img_bytes in enumerate(image_bytes_list):
        try:
            img = Image.open(io.BytesIO(img_bytes))
            if img.mode != "RGB":
                img = img.convert("RGB")

            # Resize small images for better OCR
            max_dim = max(img.size)
            if max_dim < 1500:
                scale = 1500 / max_dim
                img = img.resize((int(img.size[0]*scale), int(img.size[1]*scale)), Image.LANCZOS)

            # Enhance contrast
            img = ImageEnhance.Contrast(img).enhance(2.0)

            # Grayscale + Median filter for noise removal
            gray = img.convert("L").filter(ImageFilter.MedianFilter(size=3))

            # Binarize
            bw = gray.point(lambda x: 0 if x < 128 else 255, '1')

            # Separate OCR for French + Arabic
            text_fr = pytesseract.image_to_string(bw, lang="fra")
            text_ar = pytesseract.image_to_string(bw, lang="ara")
            text = text_fr + "\n" + text_ar

            texts.append(text)
        except Exception as e:
            print(f"[ERROR] OCR failed on image {i+1}: {e}")
            continue
    return "\n".join(texts)

# ---------- Save Text ----------
def save_text(text, rel_path, folder=CLEAN_TEXT_FOLDER):
    output_path = os.path.join(folder, rel_path + ".txt")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)
    return output_path

# ---------- Text Chunking ----------
def chunk_text(text, chunk_size=1000, overlap=200):
    """
    Split text into overlapping chunks based on sentence boundaries.
    - chunk_size: approx number of characters per chunk
    - overlap: number of overlapping characters between chunks
    """
    # Split text into sentences (works for French & Arabic)
    sentence_endings = re.compile(r'(?<=[.!?؛؟])\s+')
    sentences = sentence_endings.split(text)

    chunks = []
    current_chunk = ""
    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= chunk_size:
            current_chunk += sentence + " "
        else:
            chunks.append(current_chunk.strip())
            # Start new chunk with overlap
            overlap_text = current_chunk[-overlap:] if overlap < len(current_chunk) else current_chunk
            current_chunk = overlap_text + sentence + " "
    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks

def save_chunks(text, rel_path):
    chunks = chunk_text(text)
    base_folder = os.path.join(CHUNKS_OUTPUT, rel_path)
    os.makedirs(base_folder, exist_ok=True)
    for i, chunk in enumerate(chunks):
        chunk_rel_path = os.path.join(base_folder, f"chunk_{i+1}.txt")
        save_text(chunk, chunk_rel_path, folder="")  # folder="" because path already includes folder

# ---------- Process File ----------
def process_file(file_path, rel_path):
    if os.path.basename(file_path).startswith("~$"):
        return

    text_content = ""
    if file_path.lower().endswith(".docx"):
        text_content += extract_text_from_docx(file_path)
        images = extract_images_from_docx(file_path)
        if images:
            text_content += "\n" + ocr_images(images)
    elif file_path.lower().endswith(".pdf"):
        text_content += extract_text_from_pdf(file_path)
        images = extract_images_from_pdf(file_path)
        if images:
            text_content += "\n" + ocr_images(images)
    else:
        return

    # Clean text before saving & chunking
    text_content = clean_text(text_content)

    # Save full text
    save_text(text_content, rel_path)

    # Save chunks
    save_chunks(text_content, rel_path)

# ---------- Process All ----------
def process_all():
    for root, _, files in os.walk(RAW_FOLDER):
        for file in files:
            if file.lower().endswith((".pdf", ".docx")):
                abs_path = os.path.join(root, file)
                rel_path = os.path.splitext(os.path.relpath(abs_path, RAW_FOLDER))[0]
                process_file(abs_path, rel_path)

# ---------- Main ----------
if __name__ == "__main__":
    os.makedirs(CLEAN_TEXT_FOLDER, exist_ok=True)
    os.makedirs(CHUNKS_OUTPUT, exist_ok=True)
    process_all()
