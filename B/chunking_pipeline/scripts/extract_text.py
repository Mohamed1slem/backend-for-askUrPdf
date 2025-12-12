import os
import io
import fitz  # PyMuPDF
import docx
from PIL import Image, ImageEnhance
import pytesseract
from config import RAW_FOLDER, CLEAN_TEXT_FOLDER ,CHUNKS_OUTPUT

# Set tesseract.exe path
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# Set tessdata path
os.environ["TESSDATA_PREFIX"] = r"C:\Program Files\Tesseract-OCR\tessdata"

# ---------- Text Extraction ----------
def extract_text_from_docx(file_path):
    """Extract plain text from DOCX."""
    print(f"[DEBUG] Extracting DOCX text: {file_path}")
    doc = docx.Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    print(f"[DEBUG] Extracted {len(text)} characters from DOCX")
    return text

def extract_text_from_pdf(file_path):
    """Extract text from PDF pages."""
    print(f"[DEBUG] Extracting PDF text: {file_path}")
    doc = fitz.open(file_path)
    text = ""
    for i, page in enumerate(doc):
        page_text = page.get_text()
        print(f"[DEBUG] Page {i+1}: {len(page_text)} characters")
        text += page_text
    return text

# ---------- Image Extraction ----------
def extract_images_from_docx(file_path):
    """Extract all images from DOCX file."""
    print(f"[DEBUG] Extracting images from DOCX: {file_path}")
    doc = docx.Document(file_path)
    images = []
    for rel in doc.part._rels:
        rel_obj = doc.part._rels[rel]
        if "image" in rel_obj.target_ref:
            try:
                img_bytes = rel_obj.target_part.blob
                images.append(img_bytes)
            except Exception as e:
                print(f"[ERROR] Failed to extract DOCX image: {e}")
    print(f"[DEBUG] Found {len(images)} images in DOCX")
    return images

def extract_images_from_pdf(file_path):
    """Extract all images from PDF pages."""
    print(f"[DEBUG] Extracting images from PDF: {file_path}")
    doc = fitz.open(file_path)
    images = []
    for i, page in enumerate(doc):
        page_images = page.get_images(full=True)
        print(f"[DEBUG] Page {i+1}: {len(page_images)} images")
        for img in page_images:
            xref = img[0]
            base_image = doc.extract_image(xref)
            images.append(base_image["image"])
    print(f"[DEBUG] Total images extracted from PDF: {len(images)}")
    return images

# ---------- OCR ----------
def ocr_images(image_bytes_list):
    """Perform OCR on images with French + Arabic languages and preprocessing."""
    print(f"[DEBUG] Performing OCR on {len(image_bytes_list)} images (fra+ara)")
    texts = []

    for i, img_bytes in enumerate(image_bytes_list):
        try:
            img = Image.open(io.BytesIO(img_bytes))

            # Convert to RGB
            if img.mode != "RGB":
                img = img.convert("RGB")

            # Resize small images for better OCR
            max_dim = max(img.size)
            if max_dim < 1000:
                scale = 1000 / max_dim
                new_size = (int(img.size[0]*scale), int(img.size[1]*scale))
                img = img.resize(new_size, Image.LANCZOS)

            # Enhance contrast
            img = ImageEnhance.Contrast(img).enhance(2.0)

            # Convert to grayscale and binarize
            gray = img.convert("L")
            bw = gray.point(lambda x: 0 if x < 128 else 255, '1')

            # OCR with French + Arabic
            text = pytesseract.image_to_string(bw, lang="fra+ara")
            print(f"[DEBUG] Image {i+1}: {len(text)} characters extracted via OCR")
            texts.append(text)

        except Exception as e:
            print(f"[ERROR] Failed OCR on image {i+1}: {e}")

    return "\n".join(texts)

# ---------- Save Text ----------
def save_text(text, rel_path):
    output_path = os.path.join(CLEAN_TEXT_FOLDER, rel_path + ".txt")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"[DEBUG] Saved text to: {output_path}")

# ---------- File Processing ----------
def process_file(file_path, rel_path):
    filename = os.path.basename(file_path)
    if filename.startswith("~$"):
        print(f"[DEBUG] Skipping temp file: {filename}")
        return

    print(f"[DEBUG] Processing file: {file_path}")
    text_content = ""

    if file_path.lower().endswith(".docx"):
        text_content += extract_text_from_docx(file_path)
        images = extract_images_from_docx(file_path)
        if images:
            text_content += "\n" + ocr_images(images)

    elif file_path.lower().endswith(".pdf"):
        text_content += extract_text_from_pdf(file_path)
        images = extract_images_from_pdf(file_path)
        if images:
            text_content += "\n" + ocr_images(images)

    else:
        print(f"[DEBUG] Skipping unsupported file: {file_path}")
        return

    save_text(text_content, rel_path)
    print(f"[OK] Extracted text: {file_path} -> {rel_path}.txt\n")

# ---------- Process All ----------
def process_all():
    for root, _, files in os.walk(RAW_FOLDER):
        # Skip folders that are inside the output directories
        if CLEAN_TEXT_FOLDER in root or CHUNKS_OUTPUT in root:
            continue

        for file in files:
            if file.lower().endswith((".pdf", ".docx")):
                abs_path = os.path.join(root, file)
                rel_path = os.path.splitext(os.path.relpath(abs_path, RAW_FOLDER))[0]
                process_file(abs_path, rel_path)
# ---------- Main ----------
if __name__ == "__main__":
    os.makedirs(CLEAN_TEXT_FOLDER, exist_ok=True)
    process_all()
