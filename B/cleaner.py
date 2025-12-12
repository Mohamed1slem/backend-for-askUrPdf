import os
import re
import pdfplumber
from docx import Document
from pdf2image import convert_from_path
import pytesseract

RAW_DIR = "data/raw"
CLEAN_DIR = "data/cleaned"

def clean_text(text: str) -> str:
    """Normalize whitespace and line breaks."""
    text = re.sub(r'\s+', ' ', text)              # collapse multiple spaces
    text = re.sub(r'(\n\s*)+', '\n', text)        # normalize line breaks
    return text.strip()

def extract_keywords(text: str, top_n: int = 10):
    """Extract simple keywords by frequency for enrichment."""
    words = re.findall(r'\w+', text.lower())
    stopwords = {"the","and","for","with","this","that","les","des","dans","sur","par","من","على","في"}
    freq = {}
    for w in words:
        if w not in stopwords and len(w) > 2:
            freq[w] = freq.get(w, 0) + 1
    sorted_words = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    return [w for w, _ in sorted_words[:top_n]]

def split_into_chunks(text: str, max_len: int = 250, overlap: int = 30):
    """Split text into smaller overlapping chunks for better retrieval speed."""
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + max_len, len(words))
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        start += max_len - overlap  # overlap between chunks
    return chunks

def format_chunk(chunk: str):
    """Format chunk into Q&A style if questions are detected."""
    # Detect simple Q&A patterns
    chunk = re.sub(r"(Qui .*?\?)", r"Q: \1", chunk)
    chunk = re.sub(r"(Quel .*?\?)", r"Q: \1", chunk)
    chunk = re.sub(r"(Comment .*?\?)", r"Q: \1", chunk)
    # Add A: before answers (simple heuristic: after '?')
    chunk = re.sub(r"\?\s+", "? A: ", chunk)
    return chunk

def save_chunks(cleaned_text: str, output_path: str, filename: str):
    """Save cleaned text into multiple enriched chunked .txt files."""
    chunks = split_into_chunks(cleaned_text, max_len=250, overlap=30)
    base_name = os.path.splitext(output_path)[0]
    for i, chunk in enumerate(chunks, start=1):
        chunk = format_chunk(chunk)
        keywords = extract_keywords(chunk)
        metadata = f"[File: {filename} | Chunk: {i} | Keywords: {', '.join(keywords)}]\n"
        chunk_filename = f"{base_name}_chunk{i}.txt"
        os.makedirs(os.path.dirname(chunk_filename), exist_ok=True)
        with open(chunk_filename, "w", encoding="utf-8") as f:
            f.write(metadata + chunk + "\n\nKeywords: " + ", ".join(keywords))

def docx_to_txt(file_path: str, output_path: str):
    """Extract text and tables from DOCX and save to chunked TXT files."""
    doc = Document(file_path)
    all_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                all_text.append(" | ".join(row_text))

    raw_text = "\n".join(all_text)
    cleaned = clean_text(raw_text)
    save_chunks(cleaned, output_path, os.path.basename(file_path))

def pdf_to_txt(file_path: str, output_path: str):
    """Extract text and tables from PDF, fallback to OCR if needed, then save to chunked TXT files."""
    all_text = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                print(f"Page {i+1}: extracted {len(text)} characters")
                all_text.append(text.strip())
            else:
                print(f"Page {i+1}: no text extracted")

            # Extract tables if available
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    for row in table:
                        row_text = [cell.strip() for cell in row if cell]
                        if row_text:
                            all_text.append(" | ".join(row_text))

    # OCR fallback if no text found
    if not all_text:
        print("⚠️ No text found with pdfplumber, falling back to OCR...")
        images = convert_from_path(file_path)
        for img in images:
            text = pytesseract.image_to_string(img, lang="fra")  # use "ara" for Arabic if needed
            if text.strip():
                all_text.append(text.strip())

    raw_text = "\n".join(all_text)
    cleaned = clean_text(raw_text)
    save_chunks(cleaned, output_path, os.path.basename(file_path))

def process_all_docs():
    """Process all DOCX and PDF files in RAW_DIR recursively."""
    processed, skipped = 0, 0
    for root, _, files in os.walk(RAW_DIR):
        for filename in files:
            if filename.startswith("~$"):  # skip temp/lock files
                continue

            input_path = os.path.join(root, filename)
            rel_path = os.path.relpath(root, RAW_DIR)   # preserve subfolder structure
            output_dir = os.path.join(CLEAN_DIR, rel_path)
            os.makedirs(output_dir, exist_ok=True)

            output_filename = os.path.splitext(filename)[0] + ".txt"
            output_path = os.path.join(output_dir, output_filename)

            print(f"Processing {input_path} → {output_path}")
            try:
                if filename.lower().endswith(".docx"):
                    docx_to_txt(input_path, output_path)
                    processed += 1
                elif filename.lower().endswith(".pdf"):
                    pdf_to_txt(input_path, output_path)
                    processed += 1
                else:
                    print(f"⚠️ Skipped {input_path}: unsupported format")
                    skipped += 1
            except Exception as e:
                print(f"⚠️ Skipped {input_path}: {e}")
                skipped += 1

    print(f"✅ Done. {processed} files cleaned, {skipped} skipped.")

if __name__ == "__main__":
    process_all_docs()
